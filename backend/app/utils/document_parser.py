"""
Document URL Extractor
Extracts URLs from PDF, DOCX, TXT files
"""

import re
import io
import logging
from typing import List, Dict
from docx import Document
import PyPDF2
import pdfplumber

logger = logging.getLogger(__name__)

class DocumentParser:
    """Extract URLs from various document formats"""
    
    # URL regex pattern
    URL_PATTERN = re.compile(
        r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
    )
    
    @staticmethod
    def extract_from_pdf(file_bytes: bytes) -> Dict:
        """Extract URLs from PDF"""
        urls = []
        text_content = ""
        
        try:
            # Try pdfplumber first (better for modern PDFs)
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_content += text + "\n"
                        found_urls = DocumentParser.URL_PATTERN.findall(text)
                        for url in found_urls:
                            urls.append({
                                'url': url,
                                'page': page_num,
                                'context': DocumentParser._get_context(text, url)
                            })
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_content += text + "\n"
                        found_urls = DocumentParser.URL_PATTERN.findall(text)
                        for url in found_urls:
                            urls.append({
                                'url': url,
                                'page': page_num,
                                'context': DocumentParser._get_context(text, url)
                            })
            except Exception as e2:
                logger.error(f"PDF extraction failed: {e2}")
                raise ValueError("Unable to extract text from PDF")
        
        return {
            'urls': urls,
            'total_pages': len(text_content.split('\n')),
            'file_type': 'pdf'
        }
    
    @staticmethod
    def extract_from_docx(file_bytes: bytes) -> Dict:
        """Extract URLs from DOCX"""
        urls = []
        text_content = ""
        
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            # Extract from paragraphs
            for para_num, para in enumerate(doc.paragraphs, 1):
                text = para.text
                text_content += text + "\n"
                found_urls = DocumentParser.URL_PATTERN.findall(text)
                for url in found_urls:
                    urls.append({
                        'url': url,
                        'paragraph': para_num,
                        'context': DocumentParser._get_context(text, url)
                    })
            
            # Extract from hyperlinks
            for rel in doc.part.rels.values():
                if "hyperlink" in rel.reltype:
                    url = rel.target_ref
                    if url.startswith('http'):
                        urls.append({
                            'url': url,
                            'type': 'hyperlink',
                            'context': 'Embedded hyperlink'
                        })
        
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError("Unable to extract text from DOCX")
        
        return {
            'urls': urls,
            'total_paragraphs': len(doc.paragraphs),
            'file_type': 'docx'
        }
    
    @staticmethod
    def extract_from_txt(file_bytes: bytes) -> Dict:
        """Extract URLs from TXT"""
        try:
            text_content = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            # Try different encoding
            text_content = file_bytes.decode('latin-1')
        
        urls = []
        lines = text_content.split('\n')
        
        for line_num, line in enumerate(lines, 1):
            found_urls = DocumentParser.URL_PATTERN.findall(line)
            for url in found_urls:
                urls.append({
                    'url': url,
                    'line': line_num,
                    'context': DocumentParser._get_context(line, url)
                })
        
        return {
            'urls': urls,
            'total_lines': len(lines),
            'file_type': 'txt'
        }
    
    @staticmethod
    def _get_context(text: str, url: str, context_chars: int = 50) -> str:
        """Get surrounding context for URL"""
        try:
            idx = text.index(url)
            start = max(0, idx - context_chars)
            end = min(len(text), idx + len(url) + context_chars)
            context = text[start:end].strip()
            return context if len(context) < 150 else context[:150] + "..."
        except ValueError:
            return text[:100] + "..." if len(text) > 100 else text
    
    @staticmethod
    def extract_urls(file_bytes: bytes, file_type: str) -> Dict:
        """Main extraction method"""
        file_type = file_type.lower()
        
        if file_type == 'pdf':
            return DocumentParser.extract_from_pdf(file_bytes)
        elif file_type in ['docx', 'doc']:
            return DocumentParser.extract_from_docx(file_bytes)
        elif file_type == 'txt':
            return DocumentParser.extract_from_txt(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")